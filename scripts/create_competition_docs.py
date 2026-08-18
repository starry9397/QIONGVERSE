from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_ROOT = Path(r"D:\Lenovo\QIONVERSE琼境参赛作品")
CREATIVE_DIR = OUT_ROOT / "创作说明"
INTRO_DIR = OUT_ROOT / "作品双语介绍"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT = "F4F6F9"
PALE_BLUE = "EAF3FB"
GOLD = "A6782A"
GREEN = "22665B"
RED = "A9473D"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_header_footer(doc, short_title):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run(f"HAINAN∞QIONGVERSE 琼境  |  {short_title}")
    set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(2)
    run = fp.add_run("HAINAN∞QIONGVERSE 琼境  ·  项目材料")
    set_run_font(run, size=8.5, color=MUTED)
    fp.add_run("   ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def add_title_block(doc, kicker, title, subtitle, meta):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(kicker.upper())
    set_run_font(r, size=9, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_run_font(r, size=25, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(subtitle)
    set_run_font(r, size=12, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(meta)
    set_run_font(r, size=9.5, color=MUTED)


def add_lead(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    set_run_font(r, size=9, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.25
    r = p2.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_para(doc, text, bold_lead=None, color=INK):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, size=10.5, color=color, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, size=10.5, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=color)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, color=INK, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"资料边界：{text}")
    set_run_font(r, size=8.5, color=MUTED, italic=True)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_creative_explanation():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "创作说明")
    add_title_block(
        doc,
        "FC-OPC Next iCreate · 海南热带文旅与黎苗非遗国际出海",
        "HAINAN∞QIONGVERSE 琼境",
        "创作说明与赛事评审论证材料",
        "参赛赛道：海岛海外文旅短片、黎苗非遗数字IP、文昌航天、东方花梨、美丽乡村跨境宣传内容",
    )
    add_lead(doc, "评审摘要", "琼境把海南文化从“被观看的宣传素材”转化为“可进入、可探索、可对话、可分享的数字体验”。当前网页版本已完成六大展厅入口、沉浸式展厅、七语言、多语种 AI 导览、来源边界和海外分享路径；完整游戏化任务、收集和多人游戏闭环仍处于后续开发阶段。")

    add_section_heading(doc, "一、项目定位与参赛方向", 1)
    add_para(doc, "HAINAN∞QIONGVERSE 琼境是一项面向海外受众的海南省文化旅游数字体验项目。作品以自贸港主厅作为入口，将热带海岛、黎苗非遗、文昌航天、东方花梨和美丽乡村组织为六大沉浸式展厅，并以螺音 AI 数字人桌宠作为跨页面、跨展厅的导览角色。")
    add_para(doc, "本作品明确参赛赛道为（三）海南热带文旅与黎苗非遗国际出海：海岛海外文旅短片、黎苗非遗数字IP、文昌航天、东方花梨、美丽乡村跨境宣传内容。自贸港主厅承担跨境数字文化入口的叙事功能，五个主题展厅承担海南本土文化内容的国际化表达。")
    add_table(doc, ["项目要素", "本作品对应内容"], [
        ("核心对象", "海南省文化旅游、热带海岛、黎苗非遗、航天想象、花梨木作、乡村生活"),
        ("主要受众", "海外青年、国际游客、留学生、文化机构、城市文旅部门和跨境内容合作方"),
        ("当前交付形态", "可访问的 React/Vite 数字展览、沉浸式展厅、地图、旅行图鉴、商城演示和螺音导览"),
        ("后续形态", "持续开发中的完整游戏化探索版本，以及面向机构的定制展陈和传播服务"),
    ], [2200, 7160])
    add_source_note(doc, "赛事简章用于赛道和评分要求；当前工程蓝图用于产品愿景；网页代码、部署文档和来源登记用于当前实现事实。")

    add_section_heading(doc, "二、现实痛点与作品回应", 1)
    add_table(doc, ["传播与产业痛点", "琼境的解决方式", "当前状态"], [
        ("海外受众对海南文化缺少连续、易理解的入口", "用英文优先、七语言、地图、展厅轮盘和视觉叙事把多个文化主题串成一条可探索路径", "已实现"),
        ("传统宣传以单张图片或短视频为主，互动和记忆点不足", "用沉浸式展厅、第三人称移动、展项详情和 AI 导览把观看升级为参与", "网页体验已实现；完整游戏闭环持续开发"),
        ("非遗、航天、花梨和乡村内容分散，缺少统一品牌叙事", "以自贸港主厅和螺音为统一入口，让不同主题共享世界观、导览和分享机制", "已实现"),
        ("AI 文化内容可能出现事实幻觉和来源混淆", "使用来源登记、事实台账、项目素材标签、AIGC 概念展项标签和回答状态", "已实现"),
        ("城市和文化机构缺少可持续更新的海外数字内容工具", "将展厅、导览、多语言和社媒素材设计成可复用的数字展陈模板", "当前为项目原型，支持后续定制"),
        ("传播曝光难以形成合作线索", "通过内容分享、人工跟进意向和机构定制方向形成可控的合作入口", "人工线索为演示型、临时内存交接"),
    ], [2700, 4600, 2060], font_size=8.8)

    add_section_heading(doc, "三、核心创意与用户体验", 1)
    add_para(doc, "作品不是把多个页面简单拼接，而是设计了一条“进入海南—认识文化—与螺音对话—生成内容—分享和合作”的体验路径。用户可以从首页进入自贸港主厅，也可以直接打开地图、旅行图鉴或某一个主题展厅。")
    add_bullets(doc, [
        "自贸港主厅：提供六大展厅入口、项目定位和跨境文化入口，不替代政府办事或政策咨询。",
        "五厅轮盘：保留现有轮盘、焦点管理、键盘方向键、Tab 和进入展厅逻辑，作为视觉化文化导航。",
        "沉浸式展厅：通过 SPZ/静态回退、展项索引、详情弹层和第三人称移动，让海外用户用熟悉的探索方式接触海南文化。",
        "螺音 AI 数字人：以桌宠形态在页面和大世界中提供上下文导览；接近点位时，根据环境和展项提供多语种提示。",
        "海南地图与旅行图鉴：强调文化阅读和路线灵感，不提供实时导航、票务、预约或旅游服务承诺。",
        "商城与社交传播：商城是项目概念展示，分享功能可打开平台分享、复制链接或下载项目生成图，不虚构真实库存和交易。",
    ])

    add_section_heading(doc, "四、AI 原生完整度论证（25 分）", 1)
    add_para(doc, "琼境的 AI 原生性体现在“内容生产链”和“用户体验链”同时使用 AI，而不是只在页面中放置一个聊天框。项目将 AI 作为视觉策划、角色创作、导览生成、多语言传播和迭代验证的共同生产工具。")
    add_table(doc, ["AI 环节", "作品中的实现", "可提交的核验材料"], [
        ("世界与视觉概念", "主题展厅、场景、展项和螺音视觉的 AI 辅助创作与人工筛选", "AI 制作日志、素材目录、原始输入和衍生文件"),
        ("角色与导览", "螺音作为原创虚构 AI 数字人桌宠，支持对话、自动导览、七语言和本地回退", "导览配置、API 契约、回退文本、服务端代码和测试记录"),
        ("空间语义", "六个展厅配置导览点位，根据角色/相机姿态、展项距离和朝向触发介绍", "点位配置、触发规则、重复触发与取消请求测试"),
        ("多语言传播", "en / zh / id / ja / ko / ru / ar 内容树和 Arabic RTL 支持", "i18n 文件、语言检查脚本、七语言验收截图"),
        ("社交内容", "项目标题、固定文案、图片和视频白名单，适配 X、Facebook、TikTok、YouTube 传播路径", "分享组件、服务端资产白名单和平台状态测试"),
    ], [2100, 4500, 2760], font_size=8.8)
    add_lead(doc, "现场证据建议", "路演时展示一个展项从项目素材、AI 导览配置、来源状态到多语种分享文案的完整链路；同时打开代码或日志中的资产登记，证明 AI 生成不是不可追溯的黑箱。", fill="FFF8E8", accent=GOLD)

    add_section_heading(doc, "五、跨境合规与 AI 监管适配（20 分）", 1)
    add_para(doc, "项目在产品定位上主动限制边界：它是海南省文化旅游展示和自贸港公共信息导览层，不是政府网站、旅游预订平台、政策顾问、支付商城或官方数字人。")
    add_bullets(doc, [
        "来源分层：已核验来源、项目提供的策展素材、AIGC 策展概念展项、螺音虚构设定和 AI 建议分别标记。",
        "事实边界：不把项目图片、SPZ 场景、概念模型或 AI 文案当作真实文物、官方设施、当前旅游条件、价格、库存或政策结论。",
        "数据最小化：不保存原始移动轨迹、对话、用户画像或路线；自动导览接口只发送已登记的 cueId、展厅、语言和语音开关。",
        "密钥保护：GLM、TTS、OAuth 等密钥仅保留在服务端部署环境，不进入浏览器、构建产物、文档、日志或 URL。",
        "商业边界：商城和人工跟进只提供项目演示或意向交接，不代表订单、预订、报价、库存、合作成功或政府服务。",
        "失败透明：AI、TTS、网络、WebGL、图片、模型和分享服务失败时，保留本地文本和静态回退，不显示虚假的成功状态。",
    ])
    add_source_note(doc, "项目依据 docs/source-register.md、docs/ai-production-log.md、docs/api-safety.md 和自动导游模块说明执行内容与数据边界。")

    add_section_heading(doc, "六、赛道契合度：海南文化出海", 1)
    add_para(doc, "作品将赛道要求中的五类内容组织为一组可以被海外受众理解和传播的文化入口，而不是把它们做成互不相干的主题页面。")
    add_table(doc, ["赛道内容", "琼境中的表达", "海外传播价值"], [
        ("海岛海外文旅短片", "热带海岛厅、海岸与生态展项、旅行图鉴和项目视频", "以视觉和互动讲述海岛生态与生活方式，不依赖实时旅游承诺"),
        ("黎苗非遗数字 IP", "黎苗非遗厅、黎锦/织造语境、概念角色与螺音导览", "把纹样、手工和故事转成海外用户可理解的数字角色和展项阅读"),
        ("文昌航天", "文昌航天厅、航天概念展项和空间探索叙事", "将海南的航天想象置于文化展览框架中，避免未经核验的任务和技术承诺"),
        ("东方花梨", "东方花梨厅、木纹、雕刻、家具尺度和香器概念展项", "以材料、尺度和手工过程作为视觉语言，连接文化品牌和设计合作"),
        ("美丽乡村", "美丽乡村厅、村落、集市、路径、田野与观景语境", "表达乡村空间和生活节奏，避免把项目图像误认为实时村落信息"),
    ], [2200, 4600, 2560], font_size=8.8)

    add_section_heading(doc, "七、海外传播与商业闭环（20 分）", 1)
    add_para(doc, "琼境的商业价值不依赖夸大交易数据，而是建立一个可以被城市、展馆和文化品牌采用的内容闭环：先让受众理解和分享文化，再将明确的合作意向交给人工或机构，而不是让 AI 自动做出商业承诺。")
    add_table(doc, ["阶段", "用户/合作方动作", "项目价值与边界"], [
        ("1. 进入", "通过首页、社媒链接、城市活动或展馆二维码进入", "形成统一的海南文化数字入口；当前网页已可访问"),
        ("2. 体验", "浏览地图、轮盘、展厅、展项和螺音导览", "提升文化理解和停留时间；不收集位置轨迹"),
        ("3. 生成", "切换语言、生成项目分享内容、下载或复制链接", "把一次体验转化为可传播的多语种内容"),
        ("4. 扩散", "通过系统分享、X、Facebook 等用户控制的分享入口传播", "获得海外自然传播；TikTok/YouTube 发布需平台审核和授权配置"),
        ("5. 连接", "选择文化合作、媒体合作、工艺咨询或自贸港信息导览等人工跟进意向", "当前为最小化、同意优先、临时内存交接，不是订单或 CRM"),
        ("6. 落地", "机构采购定制展厅、城市宣传内容包或文化品牌合作", "作为后续 B2G/B2B 服务方向，需真实主体、合同、供应链和数据责任"),
    ], [1700, 3700, 3960], font_size=8.7)
    add_section_heading(doc, "八、作品应用前景", 1)
    add_bullets(doc, [
        "城市文旅国际传播：将海南城市和文化主题接入统一的数字展厅模板，服务国际活动、文旅节庆和城市宣传。",
        "博物馆与非遗展馆：将已有图片、视频、展项资料和审核来源组织为线上沉浸式展陈，降低海外访问门槛。",
        "自贸港活动数字入口：在展会、招商活动和国际会议中提供文化导览与公开信息入口，不替代正式政务渠道。",
        "高校与国际教育：作为留学生、海外高校和中华文化课程的互动案例，提供多语言、可探索的学习材料。",
        "文化品牌与设计合作：为花梨、黎锦、乡村文创和数字 IP 提供海外视觉传播、概念展陈和社交内容生产服务。",
        "后续游戏化产品：在完成完整任务、收集、成就和多人机制后，扩展为持续探索型数字文化产品。",
    ])
    add_lead(doc, "商业化前提", "当前版本不宣称真实售卖、库存、收入或合作伙伴。正式商业落地前，需要明确运营主体、授权范围、内容审核、供应链、价格、支付、售后、数据控制者和跨境平台责任。", fill="FFF1EF", accent=RED)

    add_section_heading(doc, "九、当前完成度与游戏开发计划", 1)
    add_table(doc, ["已完成或可演示", "仍在开发或规划"], [
        ("首页、六大展厅入口、五厅轮盘、海南地图、旅行图鉴、商城演示、档案馆和螺音页面", "完整开放世界游戏客户端和独立游戏启动器"),
        ("沉浸式展厅、静态回退、展项详情、鼠标/键盘/触控移动和移动端适配", "完整任务系统、神木碎片收集、成就和双结局闭环"),
        ("螺音桌宠、AI 对话、自动导览点位、多语言和 TTS/本地回退", "更完整的角色成长、可持续剧情、多人协作和正式游戏存档"),
        ("GitHub Pages、CloudBase Webify 与独立 API 的部署路径", "面向机构的正式运营后台、授权管理、长期数据治理和商业化服务"),
    ], [4680, 4680], font_size=8.8)
    add_para(doc, "因此，本次提交应把作品定位为“已完成网页数字展览与 AI 导览主体验、游戏化版本持续开发中的文化出海原型”，展示已经可以被评审体验和验证的部分，同时提出明确的下一阶段产品路线。")

    add_section_heading(doc, "十、5 分钟现场路演建议", 1)
    add_table(doc, ["时间", "展示内容", "评审记忆点"], [
        ("0:00–0:30", "痛点、赛道和一句话定位", "海南文化从静态宣传变成可探索的数字入口"),
        ("0:30–1:10", "首页、自贸港主厅和六厅轮盘", "统一叙事和清晰的文化结构"),
        ("1:10–2:00", "进入展厅，移动到展项并触发螺音", "AI 导览与空间语义结合"),
        ("2:00–2:40", "展示 AI 生产链、来源台账和回退", "AI 原生且可核验、可解释"),
        ("2:40–3:30", "切换语言、生成分享内容并打开分享选择", "海外受众路径与传播能力"),
        ("3:30–4:20", "展示商业闭环和应用场景", "从文化体验到机构合作的落地路径"),
        ("4:20–5:00", "说明当前完成度和游戏开发路线", "诚实边界、持续开发和可扩展性"),
    ], [1500, 4700, 3160], font_size=8.8)
    add_para(doc, "补充规则：赛事简章对中外学生混合组队设置额外加分。只有在团队真实符合资格并完成组委会核验时，才应在正式材料中提交相应证明；本文不虚构团队身份或加分结果。")

    add_section_heading(doc, "十一、网页预览与启动方式", 1)
    add_table(doc, ["入口", "网址/命令", "说明"], [
        ("CloudBase Webify", "https://qiongverse-webify-qiongverse-webify-d5drsw1e9a5fd4.webapps.tcloudbase.com/", "静态前端预览；通过首页导航或 Hash 路由进入各模块"),
        ("GitHub Pages", "https://starry9397.github.io/QIONGVERSE/", "静态前端预览；项目路径包含 /QIONGVERSE"),
        ("本地开发", "npm ci\nnpm run dev", "在项目根目录启动 Vite 开发服务"),
        ("本地静态预览", "npm run build\nnpm run preview", "构建后预览 dist 产物"),
        ("本地 AI 服务", ".\\Start-LuoyinGlm.ps1", "仅在本地服务进程中输入 GLM 密钥，不写入前端、文档或日志"),
    ], [1850, 4700, 2810], font_size=8.4)
    add_para(doc, "两个公网地址主要提供静态网页，AI 导览由独立 HTTPS API 服务提供。公网部署和合作方验收时，应同时检查页面资源、API 状态、CORS 白名单和隐私边界。")

    add_section_heading(doc, "十二、结语", 1)
    add_para(doc, "琼境的核心价值不是把海南文化包装成一个泛化的 AI 页面，而是建立一套可以被体验、被核验、被分享、被扩展的数字文化入口。它当前已经完成网页展览和 AI 导览主体验，下一阶段将把开放世界、收集叙事和机构服务进一步产品化，在坚持来源透明和数据最小化的前提下，为海南文化出海提供可持续的内容基础设施。")
    return doc


def bilingual_pair(doc, zh_heading, en_heading, zh, en):
    add_section_heading(doc, f"{zh_heading} / {en_heading}", 2)
    p = doc.add_paragraph()
    r = p.add_run("中文  ")
    set_run_font(r, size=9, color=GOLD, bold=True)
    r = p.add_run(zh)
    set_run_font(r, size=10.5, color=INK)
    p = doc.add_paragraph()
    r = p.add_run("English  ")
    set_run_font(r, size=9, color=GOLD, bold=True)
    r = p.add_run(en)
    set_run_font(r, size=10.5, color=INK)


def build_bilingual_introduction():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "作品双语介绍")
    add_title_block(
        doc,
        "Bilingual Project Introduction · 中英双语作品介绍",
        "HAINAN∞QIONGVERSE 琼境",
        "A living digital gateway to Hainan culture / 海南文化的可探索数字入口",
        "Track 3: Hainan tropical tourism and Li & Miao intangible heritage going global",
    )
    add_lead(doc, "Project snapshot / 项目速览", "琼境是一项面向海外受众的互动式数字文化体验。访客可以从自贸港主厅进入六大展厅，用地图、展厅轮盘、沉浸式世界和螺音 AI 数字人认识海南的海岛、黎苗非遗、文昌航天、东方花梨与美丽乡村。", fill="EAF3FB", accent=BLUE)

    bilingual_pair(doc, "项目摘要", "Project summary",
        "HAINAN∞QIONGVERSE 琼境以海南省为叙事主体，以自贸港主厅为入口，把热带海岛、黎苗非遗、文昌航天、东方花梨和美丽乡村组织为一组可探索、可对话、可分享的数字展厅。作品融合七语言界面、沉浸式展厅、AI 自动导览、海南地图、旅行图鉴、社交分享和项目概念商城，帮助海外用户用更低的理解门槛接触海南文化。",
        "HAINAN∞QIONGVERSE is an interactive digital culture experience built for international audiences. Starting from a Free Trade Port gateway, it connects tropical islands, Li & Miao heritage, Wenchang aerospace, Dongfang rosewood and beautiful villages through explorable halls, a Hainan map, a seven-language interface, social sharing and an AI guide named Luoyin."
    )

    bilingual_pair(doc, "创作理念", "Creative concept",
        "我们希望把海南文化从一组被动观看的图片，转化为一次可以进入、行走、提问和带走的数字旅程。用户不是只读一段介绍，而是以探索者身份在展厅之间移动，让文化内容通过空间、声音、角色和视觉叙事产生联系。",
        "The project turns Hainan culture from a collection of passive promotional images into a digital journey that visitors can enter, walk through, ask questions about and share. Instead of reading isolated descriptions, visitors explore as participants and connect culture through space, sound, character and visual storytelling."
    )

    add_section_heading(doc, "功能体验 / Experience", 1)
    add_table(doc, ["模块 / Module", "作品体验 / Experience", "状态 / Status"], [
        ("自贸港主厅 / Free Trade Port Hall", "六大展厅的总入口，提供项目定位和主题切换 / The gateway to six halls and the project orientation layer", "已完成 / Available"),
        ("五厅轮盘 / Hall wheel", "用视觉轮盘选择热带海岛、黎苗非遗、文昌航天、东方花梨和美丽乡村 / A visual wheel for five cultural halls", "已完成 / Available"),
        ("沉浸式展厅 / Immersive halls", "SPZ/静态回退、展项索引、详情阅读、键盘与触控移动 / SPZ or static fallback, exhibit reading, keyboard and touch movement", "已完成网页 MVP / Web MVP available"),
        ("螺音 AI 导游 / Luoyin AI guide", "桌宠、问答、空间点位触发、多语言介绍和语音回退 / Desktop pet, chat, spatial cues, multilingual guidance and speech fallback", "已完成并持续扩展 / Available and evolving"),
        ("海南地图 / Hainan map", "用项目策展语境阅读海南区域，不替代实时导航 / A project-curated regional reading layer, not live navigation", "已完成 / Available"),
        ("旅行图鉴 / Travel Atlas", "编排文化路线和视觉档案，不提供预订、票务或实时路线 / Cultural route composition without booking, tickets or live routing", "已完成 / Available"),
        ("商城与分享 / Market and sharing", "概念文创展示、复制链接、下载图像和平台分享入口 / Concept merchandise display, link copy, image download and share intents", "演示型 / Demonstration boundary"),
    ], [2100, 4760, 2200], font_size=8.6)

    bilingual_pair(doc, "螺音 AI 数字人", "Luoyin, the AI digital guide",
        "螺音是原创虚构的 AI 数字人桌宠，也是琼境的向导角色。她可以在页面中陪伴用户，在沉浸式展厅中根据展厅、展项、距离和朝向触发简短导览。她的回答会区分项目策展语境、已核验来源、AI 建议和待核验内容；当模型或网络不可用时，系统回退到本地多语言讲解。",
        "Luoyin is an original fictional AI digital guide and desktop companion. She accompanies visitors across the site and can trigger short contextual explanations when a visitor approaches an exhibit in an immersive hall. Her responses distinguish project context, reviewed sources, AI suggestions and content that needs review. When the model or network is unavailable, the experience falls back to local multilingual guidance."
    )

    bilingual_pair(doc, "五大文化主题", "Five cultural themes",
        "热带海岛厅关注海岸、生态和海岛生活；黎苗非遗厅关注织造、纹样和文化记忆；文昌航天厅将航天想象放入展览叙事；东方花梨厅阅读木纹、雕刻、家具尺度和香器概念；美丽乡村厅观察村落、田野、集市和共享空间。五个主题共同构成海南文化出海的内容骨架。",
        "The Tropical Island Hall reads coasts, ecology and island life. The Li & Miao Heritage Hall explores weaving, patterns and cultural memory. The Wenchang Aerospace Hall places aerospace imagination inside an exhibition narrative. The Dongfang Rosewood Hall focuses on grain, carving, furniture scale and scent objects. The Beautiful Villages Hall observes settlements, fields, markets and shared spaces. Together, they form the content structure for Hainan culture going global."
    )

    add_section_heading(doc, "AI 与可信内容 / AI and trusted content", 1)
    add_table(doc, ["中文", "English"], [
        ("AI 参与视觉概念、展项、角色、导览、多语言和分享内容的生产。", "AI participates in visual concepts, exhibits, character design, guidance, localization and shareable content."),
        ("项目素材不自动等于历史事实；AIGC 概念展项会明确标注。", "Project media is not automatically historical evidence; AIGC concept exhibits are explicitly labelled."),
        ("GLM、TTS 和 OAuth 密钥只在服务端使用。", "GLM, TTS and OAuth secrets remain server-side."),
        ("不保存原始移动轨迹、对话、用户画像或支付信息。", "Raw movement trails, conversations, user profiles and payment data are not stored."),
        ("商城和人工跟进是项目演示或意向交接，不是实时交易或官方服务。", "The market and human follow-up are demonstrations or intent handoffs, not live transactions or official services."),
    ], [4680, 4680], font_size=9.2)

    bilingual_pair(doc, "解决的痛点", "Problems addressed",
        "琼境回应了海外受众难以理解海南多元文化、传统宣传缺少互动、文化内容缺少统一叙事、AI 内容来源不透明以及传播难以转化为合作线索等问题。它把文化入口、互动展陈、AI 导览、多语言传播和合作意向放进同一条用户路径。",
        "The project responds to fragmented understanding of Hainan culture among international audiences, the limited interactivity of traditional promotion, the lack of a shared cultural narrative, opaque AI content provenance and the gap between exposure and qualified collaboration. It connects cultural entry, interactive exhibition, AI guidance, localization, sharing and human follow-up in one visitor path."
    )

    bilingual_pair(doc, "传播与应用前景", "Communication and application prospects",
        "琼境可以作为城市文旅国际宣传入口、博物馆和非遗展馆的线上展陈模板、自贸港活动的数字文化入口、高校国际教育案例，以及海南文化品牌和文创品牌的海外内容工具。后续可在授权、审核、运营主体和供应链明确后，扩展定制展厅、文化内容服务、IP 授权和真实商城链接。",
        "QIONGVERSE can serve as an international communication gateway for cities, an online exhibition template for museums and heritage institutions, a cultural entry point for Free Trade Port events, an educational case for international programs and a content tool for Hainan cultural brands. With proper authorization, review, operating ownership and supply-chain responsibility, it can expand into custom halls, cultural content services, IP licensing and real commerce links."
    )
    add_lead(doc, "当前版本说明 / Current release", "当前网页展览、AI 导览和多语言主体验已完成；完整开放世界游戏、任务、收集、多人协作和正式游戏存档仍处于持续开发阶段。\nThe web exhibition, AI guidance and multilingual core experience are available. The complete open-world game, quests, collection loop, multiplayer and formal save system remain under development.", fill="FFF8E8", accent=GOLD)

    add_section_heading(doc, "访问与启动 / Preview and launch", 1)
    add_table(doc, ["入口 / Entry", "地址或命令 / URL or command", "说明 / Notes"], [
        ("CloudBase Webify", "https://qiongverse-webify-qiongverse-webify-d5drsw1e9a5fd4.webapps.tcloudbase.com/", "静态前端 / Static frontend"),
        ("GitHub Pages", "https://starry9397.github.io/QIONGVERSE/", "静态前端，含项目路径 / Static frontend with project path"),
        ("Local development", "npm ci\nnpm run dev", "项目根目录执行 / Run in the project root"),
        ("Local AI service", ".\\Start-LuoyinGlm.ps1", "仅本地服务进程使用密钥 / Key stays in the local server process"),
    ], [1850, 4700, 2810], font_size=8.5)
    add_para(doc, "网页为静态前端部署，AI 导览由独立 HTTPS API 提供。The public sites serve the static frontend; AI guidance is provided by a separate HTTPS API.")

    add_section_heading(doc, "项目一句话 / One-line statement", 1)
    bilingual_pair(doc, "一句话总结", "One-line summary",
        "琼境让海外用户以探索者身份进入海南，在六大数字展厅中听见文化、看见空间，并通过螺音把一次体验带向更远的世界。",
        "QIONGVERSE lets international visitors enter Hainan as explorers, encounter culture through six digital halls, and carry the experience further through Luoyin and shareable stories."
    )
    return doc


def main():
    CREATIVE_DIR.mkdir(parents=True, exist_ok=True)
    INTRO_DIR.mkdir(parents=True, exist_ok=True)
    creative_path = CREATIVE_DIR / "HAINAN∞QIONGVERSE琼境_创作说明_终稿.docx"
    intro_path = INTRO_DIR / "HAINAN∞QIONGVERSE琼境_作品双语介绍_中英终稿.docx"
    build_creative_explanation().save(creative_path)
    build_bilingual_introduction().save(intro_path)
    print(creative_path)
    print(intro_path)


if __name__ == "__main__":
    main()
